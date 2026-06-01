from __future__ import annotations

import re
from io import BytesIO
from pathlib import Path
from typing import Tuple

from fastapi import UploadFile

from backend.services.ocr_service import OCRConfigurationError, OCRServiceError, ocr_pdf_locally
from backend.services.text_preprocessing_service import clean_for_rag

SUPPORTED_TEXT_EXTENSIONS = {".txt", ".md", ".markdown"}
SUPPORTED_DOCX_EXTENSIONS = {".docx"}
SUPPORTED_PDF_EXTENSIONS = {".pdf"}


def _decode_text(data: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "cp1256", "windows-1256", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def _default_title(filename: str, title: str | None) -> str:
    if title and title.strip():
        return title.strip()
    stem = Path(filename or "Knowledge document").stem.replace("_", " ").replace("-", " ").strip()
    return stem.title() or "Knowledge Document"


def _extract_docx_text(data: bytes) -> str:
    try:
        from docx import Document
    except Exception as exc:
        raise RuntimeError("DOCX extraction requires python-docx. Run: python3 -m pip install -r requirements.txt") from exc

    document = Document(BytesIO(data))
    chunks: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            chunks.append(text)

    # Extract tables in a readable TSV-like form.
    for table_index, table in enumerate(document.tables, start=1):
        rows: list[str] = []
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(cells):
                rows.append(" | ".join(cells))
        if rows:
            chunks.append(f"## Table {table_index}\n" + "\n".join(rows))

    return "\n\n".join(chunks).strip()


async def extract_knowledge_from_upload(file: UploadFile, title: str | None = None) -> Tuple[str, str, dict]:
    filename = file.filename or "uploaded_document"
    suffix = Path(filename).suffix.lower()
    data = await file.read()
    doc_title = _default_title(filename, title)

    if suffix in SUPPORTED_TEXT_EXTENSIONS:
        raw_content = _decode_text(data)
        source_metadata = {
            "filename": filename,
            "type": "text",
            "provider": "plain_text",
            "ocr_forced": False,
            "ocr_needed": False,
        }

    elif suffix in SUPPORTED_DOCX_EXTENSIONS:
        raw_content = _extract_docx_text(data)
        source_metadata = {
            "filename": filename,
            "type": "docx",
            "provider": "python-docx",
            "ocr_forced": False,
            "ocr_needed": False,
        }

    elif suffix in SUPPORTED_PDF_EXTENSIONS:
        # Project rule:
        # Every PDF uploaded to RAG is OCRed fully, page by page.
        # We intentionally do not use embedded PDF text extraction because table
        # layouts, Arabic text, scans, and multi-column PDFs can produce misleading text.
        try:
            result = ocr_pdf_locally(data, filename=filename)
        except OCRConfigurationError as exc:
            raise RuntimeError(str(exc)) from exc
        except OCRServiceError as exc:
            raise RuntimeError(str(exc)) from exc

        raw_content = result.text
        source_metadata = result.metadata

    else:
        raise ValueError("Unsupported file type. Please upload PDF, DOCX, TXT, MD, or Markdown files.")

    processed = clean_for_rag(raw_content)
    content = processed.text

    if not content:
        raise ValueError("No readable text could be extracted from this file.")

    metadata = {
        **source_metadata,
        "preprocessing": processed.metadata,
    }

    return doc_title, content, metadata
